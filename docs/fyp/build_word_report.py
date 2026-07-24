from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
FYP_DIR = ROOT / "docs" / "fyp"
ASSETS_DIR = FYP_DIR / "assets"
LOGO_PATH = ASSETS_DIR / "unilorin_logo_nobg.png"
OUTPUT_PATH = FYP_DIR / "Smart_Attendance_System_Final_Project_Report.docx"
FALLBACK_OUTPUT_PATH = (
    FYP_DIR / "Smart_Attendance_System_Final_Project_Report_UPDATED.docx"
)

TITLE = (
    "Design and Implementation of a Smart Attendance System Using "
    "Acoustic and Bluetooth Low Energy Proximity Verification"
)
AUTHOR_NAME = "Joshua Olugbemi Iyanuoluwa"
MATRIC_NUMBER = "21/52HP071"
REPORT_DATE = "JULY 2026"

CHAPTER_FILES = [
    FYP_DIR / "CHAPTER_ONE.md",
    FYP_DIR / "CHAPTER_TWO.md",
    FYP_DIR / "CHAPTER_THREE.md",
    FYP_DIR / "CHAPTER_FOUR.md",
    FYP_DIR / "CHAPTER_FIVE.md",
]
REFERENCES_FILE = FYP_DIR / "REFERENCES.md"


@dataclass(frozen=True)
class FigureDefinition:
    number: str
    title: str
    note: str


FIGURES = {
    "assets/report/system_architecture.png": FigureDefinition(
        "Figure 3.1",
        "Implemented System Architecture",
        "Developed by the researcher from the implemented mobile, native Android, "
        "API, and database components.",
    ),
    "assets/report/attendance_workflow.png": FigureDefinition(
        "Figure 3.2",
        "Attendance Workflow",
        "Developed by the researcher to show lecturer, student, signal, and "
        "server actions within one attendance session.",
    ),
    "assets/report/validation_flow.png": FigureDefinition(
        "Figure 3.3",
        "Backend Attendance Validation Flow",
        "Developed by the researcher from the implemented serializer, model, "
        "and database validation rules.",
    ),
    "assets/report/entity_relationship.png": FigureDefinition(
        "Figure 3.4",
        "Core Entity Relationship Model",
        "Developed by the researcher from the active logical database model.",
    ),
    "assets/report/screenshots/lecturer_session.png": FigureDefinition(
        "Figure 4.1",
        "Lecturer Session and Broadcast Interface",
        "Captured from the final installed Android application.",
    ),
    "assets/report/screenshots/student_scan.png": FigureDefinition(
        "Figure 4.2",
        "Student Scan and Proof-Review Interface",
        "Captured from the final installed Android application.",
    ),
    "assets/report/screenshots/lecturer_report.png": FigureDefinition(
        "Figure 4.3",
        "Session-Specific Lecturer Attendance Report",
        "Captured from the final installed Android application.",
    ),
    "assets/report/screenshots/beacon_admin.png": FigureDefinition(
        "Figure 4.4",
        "Registered Room-Beacon Administration Record",
        "Captured from the Django administrative interface.",
    ),
}


def _set_font(element, name: str = "Times New Roman", size: float | None = None) -> None:
    element.font.name = name
    element.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = element._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attribute}"), name)
    if size is not None:
        element.font.size = Pt(size)


def _set_style_font(style, size: float, *, bold: bool = False, italic: bool = False) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attribute}"), "Times New Roman")


def _set_outline_level(style, level: int) -> None:
    p_pr = style._element.get_or_add_pPr()
    existing = p_pr.find(qn("w:outlineLvl"))
    if existing is not None:
        p_pr.remove(existing)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    p_pr.append(outline)


def _get_or_create_style(doc: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    if name in doc.styles:
        return doc.styles[name]
    return doc.styles.add_style(name, style_type)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_style_font(normal, 12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.first_line_indent = Inches(0.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    heading_1 = doc.styles["Heading 1"]
    _set_style_font(heading_1, 12, bold=True)
    heading_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_1.paragraph_format.line_spacing = 2
    heading_1.paragraph_format.first_line_indent = Inches(0)
    heading_1.paragraph_format.space_before = Pt(12)
    heading_1.paragraph_format.space_after = Pt(0)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = doc.styles["Heading 2"]
    _set_style_font(heading_2, 12, bold=True, italic=True)
    heading_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_2.paragraph_format.line_spacing = 2
    heading_2.paragraph_format.first_line_indent = Inches(0)
    heading_2.paragraph_format.space_before = Pt(12)
    heading_2.paragraph_format.space_after = Pt(0)
    heading_2.paragraph_format.keep_with_next = True

    heading_3 = doc.styles["Heading 3"]
    _set_style_font(heading_3, 12, bold=True, italic=True)
    heading_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_3.paragraph_format.line_spacing = 2
    heading_3.paragraph_format.first_line_indent = Inches(0)
    heading_3.paragraph_format.space_before = Pt(8)
    heading_3.paragraph_format.space_after = Pt(0)
    heading_3.paragraph_format.keep_with_next = True

    chapter = _get_or_create_style(doc, "Chapter Title")
    _set_style_font(chapter, 14, bold=True)
    chapter.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chapter.paragraph_format.line_spacing = 2
    chapter.paragraph_format.first_line_indent = Inches(0)
    chapter.paragraph_format.space_before = Pt(0)
    chapter.paragraph_format.space_after = Pt(12)
    chapter.paragraph_format.keep_with_next = True
    _set_outline_level(chapter, 0)

    preliminary = _get_or_create_style(doc, "Preliminary Heading")
    _set_style_font(preliminary, 14, bold=True)
    preliminary.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    preliminary.paragraph_format.line_spacing = 2
    preliminary.paragraph_format.first_line_indent = Inches(0)
    preliminary.paragraph_format.space_before = Pt(0)
    preliminary.paragraph_format.space_after = Pt(18)
    preliminary.paragraph_format.keep_with_next = True
    _set_outline_level(preliminary, 0)

    no_toc_heading = _get_or_create_style(doc, "Front Heading No TOC")
    _set_style_font(no_toc_heading, 14, bold=True)
    no_toc_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_toc_heading.paragraph_format.line_spacing = 2
    no_toc_heading.paragraph_format.first_line_indent = Inches(0)
    no_toc_heading.paragraph_format.space_after = Pt(18)

    list_item = _get_or_create_style(doc, "Report List")
    _set_style_font(list_item, 12)
    list_item.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    list_item.paragraph_format.line_spacing = 2
    list_item.paragraph_format.left_indent = Inches(0.5)
    list_item.paragraph_format.first_line_indent = Inches(-0.25)
    list_item.paragraph_format.space_before = Pt(0)
    list_item.paragraph_format.space_after = Pt(0)

    reference = _get_or_create_style(doc, "APA Reference")
    _set_style_font(reference, 12)
    reference.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    reference.paragraph_format.line_spacing = 2
    reference.paragraph_format.left_indent = Inches(0.5)
    reference.paragraph_format.first_line_indent = Inches(-0.5)
    reference.paragraph_format.space_before = Pt(0)
    reference.paragraph_format.space_after = Pt(0)

    figure_caption = _get_or_create_style(doc, "Figure Caption")
    _set_style_font(figure_caption, 11)
    figure_caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    figure_caption.paragraph_format.line_spacing = 1
    figure_caption.paragraph_format.first_line_indent = Inches(0)
    figure_caption.paragraph_format.space_before = Pt(8)
    figure_caption.paragraph_format.space_after = Pt(5)
    figure_caption.paragraph_format.keep_with_next = True

    table_caption = _get_or_create_style(doc, "Table Caption")
    _set_style_font(table_caption, 11)
    table_caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table_caption.paragraph_format.line_spacing = 1
    table_caption.paragraph_format.first_line_indent = Inches(0)
    table_caption.paragraph_format.space_before = Pt(8)
    table_caption.paragraph_format.space_after = Pt(5)
    table_caption.paragraph_format.keep_with_next = True

    figure_note = _get_or_create_style(doc, "Figure Note")
    _set_style_font(figure_note, 10)
    figure_note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    figure_note.paragraph_format.line_spacing = 1
    figure_note.paragraph_format.first_line_indent = Inches(0)
    figure_note.paragraph_format.space_before = Pt(4)
    figure_note.paragraph_format.space_after = Pt(8)

    code = _get_or_create_style(doc, "Report Code")
    _set_style_font(code, 10)
    code.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    code.paragraph_format.line_spacing = 1
    code.paragraph_format.left_indent = Inches(0.35)
    code.paragraph_format.first_line_indent = Inches(0)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)


def configure_page(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def _field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    if placeholder:
        text = OxmlElement("w:t")
        text.text = placeholder
        run._r.append(text)
    run._r.append(end)
    _set_font(run, size=12)


def _set_page_number_format(section, *, fmt: str, start: int) -> None:
    sect_pr = section._sectPr
    page_number = sect_pr.find(qn("w:pgNumType"))
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        sect_pr.append(page_number)
    page_number.set(qn("w:fmt"), fmt)
    page_number.set(qn("w:start"), str(start))


def _set_footer_page_number(section) -> None:
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(paragraph, "PAGE")


def _disable_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    _clear_paragraph(paragraph)


def _request_field_updates(doc: Document) -> None:
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_formatted_runs(paragraph, text: str) -> None:
    for part in INLINE_PATTERN.split(text):
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            _set_font(run, size=10)
            continue
        else:
            run.text = part
        _set_font(run, size=12)


def add_body_paragraph(
    doc: Document,
    text: str,
    *,
    first_line_indent: bool = True,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = alignment
    if not first_line_indent:
        paragraph.paragraph_format.first_line_indent = Inches(0)
    add_formatted_runs(paragraph, text)
    return paragraph


def add_manual_list_item(doc: Document, marker: str, text: str) -> None:
    paragraph = doc.add_paragraph(style="Report List")
    marker_run = paragraph.add_run(f"{marker} ")
    _set_font(marker_run, size=12)
    add_formatted_runs(paragraph, text)


def _set_cell_margins(cell, *, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge, value, size in (
        ("top", "single", "12"),
        ("bottom", "single", "12"),
        ("insideH", "single", "6"),
        ("left", "nil", "0"),
        ("right", "nil", "0"),
        ("insideV", "nil", "0"),
    ):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), value)
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")
        borders.append(node)
    tbl_pr.append(borders)


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _clean_markdown(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return re.sub(r"`([^`]+)`", r"\1", text)


def add_table_caption(doc: Document, number: str, title: str) -> None:
    paragraph = doc.add_paragraph(style="Table Caption")
    number_run = paragraph.add_run(number)
    number_run.bold = True
    _set_font(number_run, size=11)
    number_run.add_break()
    title_run = paragraph.add_run(f" {title}")
    title_run.italic = True
    _set_font(title_run, size=11)


def add_markdown_table(
    doc: Document,
    rows: list[list[str]],
    caption: tuple[str, str] | None,
) -> None:
    if not rows:
        return
    if caption is not None:
        add_table_caption(doc, *caption)

    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders(table)
    _repeat_table_header(table.rows[0])

    cell_font_size = 9 if column_count >= 4 else 10
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_margins(cell)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Inches(0)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            text = row[column_index] if column_index < len(row) else ""
            run = paragraph.add_run(_clean_markdown(text))
            run.bold = row_index == 0
            _set_font(run, size=cell_font_size)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 1
    spacer.add_run("")


def add_figure_caption(doc: Document, definition: FigureDefinition) -> None:
    paragraph = doc.add_paragraph(style="Figure Caption")
    number_run = paragraph.add_run(definition.number)
    number_run.bold = True
    _set_font(number_run, size=11)
    number_run.add_break()
    title_run = paragraph.add_run(f" {definition.title}")
    title_run.italic = True
    _set_font(title_run, size=11)


def add_figure(doc: Document, path: Path, definition: FigureDefinition) -> None:
    add_figure_caption(doc, definition)
    with Image.open(path) as image:
        width, height = image.size
    aspect_ratio = width / max(height, 1)
    width_inches = min(6.25, 7.5 * aspect_ratio)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_inches))

    note = doc.add_paragraph(style="Figure Note")
    prefix = note.add_run("Note. ")
    prefix.italic = True
    _set_font(prefix, size=10)
    body = note.add_run(definition.note)
    _set_font(body, size=10)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    for run in paragraph.runs:
        _set_font(run, size=12)


def add_chapter_title(doc: Document, markdown_title: str) -> None:
    parts = markdown_title.split(":", maxsplit=1)
    paragraph = doc.add_paragraph(style="Chapter Title")
    first = parts[0].strip().upper()
    second = parts[1].strip().upper() if len(parts) == 2 else ""
    run = paragraph.add_run(first)
    run.bold = True
    _set_font(run, size=14)
    if second:
        run.add_break()
        second_run = paragraph.add_run(second)
        second_run.bold = True
        _set_font(second_run, size=14)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    separator = re.compile(
        r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$"
    )
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip()
        if not separator.match(raw):
            rows.append([cell.strip() for cell in raw.strip("|").split("|")])
        index += 1
    return rows, index


TABLE_CAPTION_PATTERN = re.compile(
    r"^\*\*(Table\s+\d+\.\d+)\s*:\s*(.+?)\*\*$",
    flags=re.IGNORECASE,
)
IMAGE_PATTERN = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
IMAGE_CAPTION_PATTERN = re.compile(r"^\*Figure\s+\d+\.\d+\..+\*$", re.IGNORECASE)


def parse_markdown_file(
    doc: Document,
    path: Path,
    *,
    is_references: bool = False,
) -> list[str]:
    lines = path.read_text(encoding="utf-8").splitlines()
    missing_images: list[str] = []
    pending_table_caption: tuple[str, str] | None = None
    code_lines: list[str] = []
    in_code = False
    index = 0

    while index < len(lines):
        raw_line = lines[index].rstrip()
        stripped = raw_line.strip()

        if stripped.startswith("```"):
            if in_code:
                for code_line in code_lines:
                    paragraph = doc.add_paragraph(style="Report Code")
                    run = paragraph.add_run(code_line)
                    _set_font(run, size=10)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(raw_line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        table_caption_match = TABLE_CAPTION_PATTERN.match(stripped)
        if table_caption_match:
            pending_table_caption = (
                table_caption_match.group(1).title(),
                table_caption_match.group(2).rstrip("."),
            )
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows, pending_table_caption)
            pending_table_caption = None
            continue

        image_match = IMAGE_PATTERN.match(stripped)
        if image_match:
            relative_path = image_match.group(2).replace("\\", "/")
            definition = FIGURES.get(relative_path)
            image_path = FYP_DIR / relative_path
            if definition is not None and image_path.exists():
                add_figure(doc, image_path, definition)
            else:
                missing_images.append(relative_path)
            index += 1
            continue

        if IMAGE_CAPTION_PATTERN.match(stripped):
            index += 1
            continue

        if stripped.startswith("# "):
            heading = stripped[2:].strip()
            if is_references:
                add_chapter_title(doc, "References")
            else:
                add_chapter_title(doc, heading)
            index += 1
            continue

        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 1)
            index += 1
            continue

        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 2)
            index += 1
            continue

        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:].strip(), 3)
            index += 1
            continue

        if stripped.startswith(">"):
            paragraph = add_body_paragraph(
                doc,
                stripped.lstrip("> ").strip(),
                first_line_indent=False,
            )
            paragraph.paragraph_format.left_indent = Inches(0.5)
            for run in paragraph.runs:
                run.italic = True
            index += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            add_manual_list_item(doc, "\u2022", bullet.group(1))
            index += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            add_manual_list_item(doc, f"{numbered.group(1)}.", numbered.group(2))
            index += 1
            continue

        if is_references:
            paragraph = doc.add_paragraph(style="APA Reference")
            add_formatted_runs(paragraph, stripped)
        else:
            add_body_paragraph(doc, stripped)
        index += 1

    return missing_images


def add_centered_line(
    doc: Document,
    text: str,
    *,
    size: float = 12,
    bold: bool = False,
    space_before: float = 0,
    space_after: float = 0,
):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.bold = bold
    _set_font(run, size=size)
    return paragraph


def add_title_page(doc: Document) -> None:
    if LOGO_PATH.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(1.35))

    add_centered_line(doc, TITLE.upper(), size=14, bold=True, space_after=18)
    add_centered_line(doc, "BY", bold=True, space_after=6)
    add_centered_line(doc, AUTHOR_NAME.upper(), bold=True)
    add_centered_line(doc, MATRIC_NUMBER, bold=True, space_after=20)
    add_centered_line(
        doc,
        "A FINAL YEAR PROJECT SUBMITTED TO THE DEPARTMENT OF "
        "TELECOMMUNICATION SCIENCE,",
        size=12,
    )
    add_centered_line(
        doc,
        "FACULTY OF COMMUNICATION AND INFORMATION SCIENCES,",
        size=12,
    )
    add_centered_line(
        doc,
        "UNIVERSITY OF ILORIN, ILORIN, NIGERIA",
        size=12,
        space_after=18,
    )
    add_centered_line(
        doc,
        "IN PARTIAL FULFILMENT OF THE REQUIREMENTS FOR THE AWARD OF "
        "THE DEGREE OF BACHELOR OF SCIENCE (B.Sc.) IN "
        "TELECOMMUNICATION SCIENCE",
        size=12,
        space_after=22,
    )
    add_centered_line(doc, REPORT_DATE, bold=True)


def add_front_heading(doc: Document, title: str, *, include_in_toc: bool = True) -> None:
    style = "Preliminary Heading" if include_in_toc else "Front Heading No TOC"
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(title.upper())
    run.bold = True
    _set_font(run, size=14)


def add_signature_table(doc: Document, labels: list[str], columns: int = 2) -> None:
    rows = (len(labels) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, label in enumerate(labels):
        row = index // columns
        column = index % columns
        cell = table.cell(row, column)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.add_run("\n____________________________\n")
        label_run = paragraph.add_run(label)
        label_run.bold = True
        date_run = paragraph.add_run("\nSignature and Date")
        for run in paragraph.runs:
            _set_font(run, size=11)
        _set_cell_margins(cell, top=80, bottom=80, start=120, end=120)
    _set_table_borders_none(table)


def _set_table_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)


def add_declaration(doc: Document) -> None:
    add_front_heading(doc, "Declaration")
    body = (
        f"I, {AUTHOR_NAME}, with matriculation number {MATRIC_NUMBER}, declare "
        f"that this project titled \u201c{TITLE}\u201d is my original work, carried "
        "out in the Department of Telecommunication Science, University of "
        "Ilorin. All ideas, findings, and words obtained from other authors have "
        "been acknowledged through appropriate citation and referencing. This "
        "work has not been submitted, either wholly or partly, to this or any "
        "other institution for the award of a degree or qualification."
    )
    add_body_paragraph(doc, body)
    add_signature_table(doc, [AUTHOR_NAME], columns=1)


def add_certification(doc: Document) -> None:
    add_front_heading(doc, "Certification")
    body = (
        f"This is to certify that the project titled \u201c{TITLE}\u201d was "
        f"carried out by {AUTHOR_NAME} ({MATRIC_NUMBER}) in the Department of "
        "Telecommunication Science, University of Ilorin, under approved "
        "academic supervision. The project has been examined and found to meet "
        "the requirements for the award of the Bachelor of Science (B.Sc.) "
        "degree in Telecommunication Science."
    )
    add_body_paragraph(doc, body)
    add_signature_table(
        doc,
        [
            "Project Supervisor",
            "Head of Department",
            "External Examiner",
        ],
    )


def add_copyright(doc: Document) -> None:
    add_front_heading(doc, "Copyright")
    add_body_paragraph(
        doc,
        f"\u00a9 2026 {AUTHOR_NAME}. All rights reserved.",
        first_line_indent=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_body_paragraph(
        doc,
        "No part of this project report may be reproduced, stored in a retrieval "
        "system, or transmitted in any form or by any means without prior "
        "written permission from the author, except for brief quotations used "
        "for academic review, criticism, or other purposes permitted by law.",
    )


def add_dedication(doc: Document) -> None:
    add_front_heading(doc, "Dedication")
    paragraph = add_body_paragraph(
        doc,
        "This work is dedicated to God Almighty, whose grace and strength "
        "sustained me throughout the course of this study; to the cherished "
        "memory of my late father; and to my family, whose love, sacrifice, and "
        "encouragement have remained a constant source of support.",
        first_line_indent=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    paragraph.paragraph_format.space_before = Pt(80)


def add_acknowledgements(doc: Document) -> None:
    add_front_heading(doc, "Acknowledgements")
    paragraphs = [
        (
            "I give all glory to God Almighty for the wisdom, strength, and "
            "opportunity to complete this project."
        ),
        (
            "I sincerely appreciate my project supervisor for the guidance, "
            "constructive observations, and academic direction provided during "
            "the development of this work. I am also grateful to the academic "
            "and technical staff of the Department of Telecommunication "
            "Science, University of Ilorin, for the knowledge and support that "
            "contributed to my training."
        ),
        (
            "My profound gratitude goes to my family for their patience, "
            "sacrifices, prayers, and unwavering encouragement. I equally thank "
            "my colleagues and friends whose discussions, practical assistance, "
            "and willingness to participate in testing helped the project "
            "progress from an idea to a working implementation."
        ),
    ]
    for paragraph in paragraphs:
        add_body_paragraph(doc, paragraph)


def add_abstract(doc: Document) -> None:
    add_front_heading(doc, "Abstract")
    abstract = (
        "This study designed and implemented a mobile attendance system that "
        "uses acoustic beaconing and Bluetooth Low Energy (BLE) proximity "
        "evidence to support classroom attendance. A Flutter application "
        "provides separate lecturer and student workflows, while native Kotlin "
        "components perform acoustic transmission and decoding, BLE advertising "
        "and scanning, and foreground broadcast management. A Django REST API "
        "deployed on Render validates authenticated submissions and stores "
        "accepted records in a relational database. The validation process "
        "checks the selected session, signal format and age, registered "
        "beacon-room association, received signal strength policy, proof "
        "digest, student identity, device binding, duplicate attendance, and "
        "student-scoped replay. Development followed an iterative Agile "
        "approach. Verification comprised 15 backend API tests, four Flutter "
        "tests, Android compilation, and repeated small-scale physical testing "
        "with four Android phones and a DX-CP27 Mini beacon. The deterministic "
        "software checks passed. Acoustic decoding was normally observed at "
        "approximately 1\u201330 cm and slightly beyond 50 cm in the best quiet, "
        "high-volume condition. Lecturer-device BLE was detected by the three "
        "student phones at 1\u20135 m, 5\u201310 m, and above 10 m under the tested "
        "conditions, including some indoor obstructions. The room beacon was "
        "also detected and resolved through its registered classroom. These "
        "results show that BLE is the practical primary path in the current "
        "prototype, whereas acoustic proof provides supplementary very-short-"
        "range evidence. The evaluation was limited by estimated distances, "
        "the absence of fixed trial counts, four phones, and no full-class or "
        "adjacent-room test; consequently, the findings establish "
        "prototype feasibility rather than institution-wide performance."
    )
    add_body_paragraph(doc, abstract, first_line_indent=False)
    keywords = doc.add_paragraph()
    keywords.paragraph_format.first_line_indent = Inches(0)
    keywords.paragraph_format.line_spacing = 2
    prefix = keywords.add_run("Keywords: ")
    prefix.italic = True
    _set_font(prefix, size=12)
    body = keywords.add_run(
        "smart attendance, acoustic beacon, Bluetooth Low Energy, proximity "
        "verification, room beacon, device binding"
    )
    _set_font(body, size=12)


def add_contents_page(doc: Document) -> None:
    add_front_heading(doc, "Table of Contents", include_in_toc=False)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0)
    _field(
        paragraph,
        r'TOC \o "1-3" \h \z \u',
        "Open this document in Microsoft Word and update the field.",
    )


def add_list_page(doc: Document, title: str, caption_style: str) -> None:
    add_front_heading(doc, title)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0)
    _field(
        paragraph,
        rf'TOC \h \z \t "{caption_style},1"',
        "Open this document in Microsoft Word and update the field.",
    )


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_preliminary_pages(doc: Document) -> None:
    page_builders = [
        add_declaration,
        add_certification,
        add_copyright,
        add_dedication,
        add_acknowledgements,
        add_abstract,
        add_contents_page,
        lambda target: add_list_page(target, "List of Figures", "Figure Caption"),
        lambda target: add_list_page(target, "List of Tables", "Table Caption"),
    ]
    for index, builder in enumerate(page_builders):
        builder(doc)
        if index < len(page_builders) - 1:
            add_page_break(doc)


def enforce_times_new_roman(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            current_size = run.font.size.pt if run.font.size is not None else None
            _set_font(run, size=current_size)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        current_size = (
                            run.font.size.pt if run.font.size is not None else None
                        )
                        _set_font(run, size=current_size)
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                for run in paragraph.runs:
                    current_size = (
                        run.font.size.pt if run.font.size is not None else None
                    )
                    _set_font(run, size=current_size)


def build_report() -> Path:
    doc = Document()
    configure_styles(doc)
    configure_page(doc.sections[0])
    _disable_footer(doc.sections[0])
    _request_field_updates(doc)

    properties = doc.core_properties
    properties.title = TITLE
    properties.subject = "Final Year Project Report"
    properties.author = AUTHOR_NAME
    properties.last_modified_by = AUTHOR_NAME
    properties.keywords = (
        "smart attendance; acoustic beacon; Bluetooth Low Energy; "
        "proximity verification"
    )
    properties.comments = "Prepared from the verified project report sources."

    add_title_page(doc)

    front_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(front_section)
    _set_page_number_format(front_section, fmt="lowerRoman", start=1)
    _set_footer_page_number(front_section)
    add_preliminary_pages(doc)

    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(main_section)
    _set_page_number_format(main_section, fmt="decimal", start=1)
    _set_footer_page_number(main_section)

    missing_images: list[str] = []
    for index, chapter_path in enumerate(CHAPTER_FILES):
        if index > 0:
            add_page_break(doc)
        missing_images.extend(parse_markdown_file(doc, chapter_path))

    add_page_break(doc)
    missing_images.extend(
        parse_markdown_file(doc, REFERENCES_FILE, is_references=True)
    )

    enforce_times_new_roman(doc)

    output = OUTPUT_PATH
    try:
        doc.save(output)
    except PermissionError:
        output = FALLBACK_OUTPUT_PATH
        doc.save(output)

    print(f"Generated: {output}")
    if missing_images:
        print("Skipped unavailable implementation screenshots:")
        for item in sorted(set(missing_images)):
            print(f"  - {item}")
    return output


if __name__ == "__main__":
    build_report()
