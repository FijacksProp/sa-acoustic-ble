from __future__ import annotations

import argparse
import json
import re
import zipfile
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.oxml.ns import qn


def _font_name(run) -> str:
    if run.font.name:
        return run.font.name
    r_pr = run._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return ""
    return (
        r_pr.rFonts.get(qn("w:ascii"))
        or r_pr.rFonts.get(qn("w:hAnsi"))
        or ""
    )


def _all_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def audit(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip()]
    all_paragraphs = [
        paragraph for paragraph in _all_paragraphs(doc) if paragraph.text.strip()
    ]
    heading_counts = Counter(
        paragraph.style.name
        for paragraph in paragraphs
        if paragraph.style.name.startswith("Heading")
    )
    run_fonts = Counter()
    non_times_runs = []
    for paragraph in all_paragraphs:
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            font = _font_name(run) or "(inherited)"
            run_fonts[font] += 1
            if font not in {"Times New Roman", "(inherited)"}:
                non_times_runs.append(
                    {
                        "font": font,
                        "text": run.text.strip()[:120],
                    }
                )

    suspicious_terms = re.compile(
        r"\b(?:draft|placeholder|insert figure|write your|"
        r"to be completed|should be updated)\b",
        re.IGNORECASE,
    )
    suspicious_paragraphs = [
        paragraph.text.strip()
        for paragraph in paragraphs
        if suspicious_terms.search(paragraph.text)
    ]
    numbered_paragraphs = [
        paragraph.text.strip()
        for paragraph in paragraphs
        if re.match(r"^\d+[.)]\s+", paragraph.text.strip())
    ]

    sections = []
    for section in doc.sections:
        sections.append(
            {
                "page_width_inches": round(section.page_width.inches, 3),
                "page_height_inches": round(section.page_height.inches, 3),
                "top_margin_inches": round(section.top_margin.inches, 3),
                "bottom_margin_inches": round(section.bottom_margin.inches, 3),
                "left_margin_inches": round(section.left_margin.inches, 3),
                "right_margin_inches": round(section.right_margin.inches, 3),
                "different_first_page": section.different_first_page_header_footer,
            }
        )

    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        document_xml = archive.read("word/document.xml").decode(
            "utf-8",
            errors="replace",
        )
        footer_xml = "".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in names
            if name.startswith("word/footer") and name.endswith(".xml")
        )
        footnotes_xml = (
            archive.read("word/footnotes.xml").decode("utf-8", errors="replace")
            if "word/footnotes.xml" in names
            else ""
        )
        endnotes_xml = (
            archive.read("word/endnotes.xml").decode("utf-8", errors="replace")
            if "word/endnotes.xml" in names
            else ""
        )
        comments_xml = (
            archive.read("word/comments.xml").decode("utf-8", errors="replace")
            if "word/comments.xml" in names
            else ""
        )

    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def count_authored_notes(xml: str, element_name: str) -> int:
        if not xml:
            return 0
        root = ElementTree.fromstring(xml)
        excluded_types = {"separator", "continuationSeparator", "continuationNotice"}
        return sum(
            note.get(f"{{{namespace}}}type") not in excluded_types
            for note in root.findall(f"{{{namespace}}}{element_name}")
        )

    footnote_count = count_authored_notes(footnotes_xml, "footnote")
    endnote_count = count_authored_notes(endnotes_xml, "endnote")
    comment_count = len(re.findall(r"<w:comment(?:\s|>)", comments_xml))

    return {
        "path": str(path),
        "file_size_bytes": path.stat().st_size,
        "paragraph_count": len(paragraphs),
        "all_container_paragraph_count": len(all_paragraphs),
        "table_count": len(doc.tables),
        "table_shapes": [
            [len(table.rows), len(table.columns)]
            for table in doc.tables
        ],
        "inline_image_count": len(doc.inline_shapes),
        "section_count": len(doc.sections),
        "sections": sections,
        "heading_counts": dict(heading_counts),
        "heading_text": [
            paragraph.text.strip()
            for paragraph in paragraphs
            if paragraph.style.name.startswith("Heading")
        ],
        "run_fonts": dict(run_fonts),
        "non_times_new_roman_runs": non_times_runs[:50],
        "suspicious_paragraphs": suspicious_paragraphs,
        "manual_numbered_paragraph_count": len(numbered_paragraphs),
        "manual_numbered_paragraphs": numbered_paragraphs[:40],
        "has_toc_field": 'w:instrText xml:space="preserve">TOC ' in document_xml,
        "has_page_field": bool(re.search(r">\s*PAGE\s*<", footer_xml)),
        "footnote_count": footnote_count,
        "endnote_count": endnote_count,
        "comment_count": comment_count,
        "has_footnotes": footnote_count > 0,
        "has_endnotes": endnote_count > 0,
        "has_comments": comment_count > 0,
        "has_tracked_insertions": bool(
            re.search(r"<w:ins(?:\s|>)", document_xml)
        ),
        "has_tracked_deletions": bool(
            re.search(r"<w:del(?:\s|>)", document_xml)
        ),
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("document", type=Path)
    parser.add_argument("--summary", action="store_true")
    args = parser.parse_args()
    result = audit(args.document)
    if args.summary:
        result = {
            "path": result["path"],
            "file_size_bytes": result["file_size_bytes"],
            "paragraph_count": result["paragraph_count"],
            "all_container_paragraph_count": result[
                "all_container_paragraph_count"
            ],
            "table_count": result["table_count"],
            "table_shapes": result["table_shapes"],
            "inline_image_count": result["inline_image_count"],
            "section_count": result["section_count"],
            "sections": result["sections"],
            "heading_counts": result["heading_counts"],
            "heading_text": result["heading_text"],
            "run_fonts": result["run_fonts"],
            "non_times_new_roman_run_count": len(
                result["non_times_new_roman_runs"]
            ),
            "non_times_new_roman_runs": result[
                "non_times_new_roman_runs"
            ][:10],
            "suspicious_paragraph_count": len(
                result["suspicious_paragraphs"]
            ),
            "suspicious_paragraphs": result[
                "suspicious_paragraphs"
            ][:20],
            "manual_numbered_paragraph_count": result[
                "manual_numbered_paragraph_count"
            ],
            "manual_numbered_paragraphs": result[
                "manual_numbered_paragraphs"
            ][:10],
            "has_toc_field": result["has_toc_field"],
            "has_page_field": result["has_page_field"],
            "footnote_count": result["footnote_count"],
            "endnote_count": result["endnote_count"],
            "comment_count": result["comment_count"],
            "has_footnotes": result["has_footnotes"],
            "has_endnotes": result["has_endnotes"],
            "has_comments": result["has_comments"],
            "has_tracked_insertions": result[
                "has_tracked_insertions"
            ],
            "has_tracked_deletions": result[
                "has_tracked_deletions"
            ],
        }
    print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
